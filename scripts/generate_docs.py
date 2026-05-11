"""Generate professional Voice Outbound Agent technical documentation."""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_BODY = "Calibri"
FONT_CODE = "Courier New"

def _shd(fill):
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear"); e.set(qn("w:color"), "auto"); e.set(qn("w:fill"), fill)
    return e

def _border_bottom(p, color="000000", sz="8"):
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
    pBdr.append(b); pPr.append(pBdr)

def _run(para, text, size=11, bold=False, italic=False, name=None):
    r = para.add_run(text)
    r.font.name = name or FONT_BODY
    r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    return r

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    _run(p, text, size=20, bold=True)
    _border_bottom(p)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    _run(p, text, size=14, bold=True)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _run(p, text, size=12, bold=True, italic=True)
    return p

def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    _run(p, text)
    return p

def code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, line if line.strip() else " ", size=9, name=FONT_CODE)
        p._element.get_or_add_pPr().append(_shd("F2F2F2"))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def diag(doc, text, caption=None):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, line if line.strip() else " ", size=8, name=FONT_CODE)
        p._element.get_or_add_pPr().append(_shd("F8F8F8"))
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(10)
        _run(cp, caption, size=9, italic=True)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        for pr in hcells[i].paragraphs:
            for r in pr.runs:
                r.font.name = FONT_BODY; r.font.size = Pt(10); r.bold = True
        hcells[i]._element.get_or_add_tcPr().append(_shd("D9D9D9"))
    for ri, row_data in enumerate(rows):
        rc = t.rows[ri + 1].cells
        for ci, txt in enumerate(row_data):
            rc[ci].text = txt
            for pr in rc[ci].paragraphs:
                for r in pr.runs:
                    r.font.name = FONT_BODY; r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def bul(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    _run(p, text)

def num(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    _run(p, text)

def pb(doc):
    doc.add_page_break()


def build(output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Title Page ──────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(tp, "VOICE OUTBOUND AGENT", size=32, bold=True)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sp, "System Architecture and Technical Reference", size=16)

    doc.add_paragraph()
    div = doc.add_paragraph()
    _border_bottom(div)

    for _ in range(2):
        doc.add_paragraph()

    for line in [
        "Prepared by: Fidelitus Corp / SherpaVector",
        "Document Version: 1.0",
        "Date: May 2026",
    ]:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(lp, line, size=11)

    for _ in range(4):
        doc.add_paragraph()

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cp, "CONFIDENTIAL — For internal use only", size=9, italic=True)

    pb(doc)

    # ── Document History ────────────────────────────────────────────────────
    h1(doc, "Document History")
    tbl(doc,
        ["Version", "Date", "Author", "Changes"],
        [["1.0", "May 2026", "Fidelitus Corp / SherpaVector", "Initial release — all 11 modules complete"]]
    )

    # ── Table of Contents ───────────────────────────────────────────────────
    h1(doc, "Table of Contents")
    toc = [
        ("1.", "Introduction", True),
        ("  1.1", "Purpose of This Document", False),
        ("  1.2", "Scope and Audience", False),
        ("  1.3", "Document Conventions", False),
        ("2.", "System Architecture", True),
        ("  2.1", "High-Level Overview", False),
        ("  2.2", "Technology Stack", False),
        ("  2.3", "Service Topology", False),
        ("  2.4", "Database Schema", False),
        ("  2.5", "Security Architecture", False),
        ("3.", "Data Flow Diagrams", True),
        ("  3.1", "Outbound Call Initiation", False),
        ("  3.2", "Webhook Event Processing", False),
        ("  3.3", "Post-Call Analysis Pipeline", False),
        ("  3.4", "DNC Enforcement Flow", False),
        ("4.", "Local Development Setup", True),
        ("  4.1", "Prerequisites", False),
        ("  4.2", "Repository Setup", False),
        ("  4.3", "Environment Configuration", False),
        ("  4.4", "Starting the Full Stack", False),
        ("  4.5", "Verification and Health Checks", False),
        ("  4.6", "Running the Test Suite", False),
        ("5.", "Module Reference", True),
        ("6.", "API Reference", True),
        ("7.", "Deployment Guide", True),
        ("8.", "Compliance and Operations", True),
        ("Appendix A", "Environment Variables Reference", True),
        ("Appendix B", "Glossary", True),
    ]
    for num_str, title, bold in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, f"{num_str}   {title}", size=11, bold=bold)

    pb(doc)

    # ── Section 1: Introduction ─────────────────────────────────────────────
    h1(doc, "1.  Introduction")

    h2(doc, "1.1  Purpose of This Document")
    para(doc,
        "This document provides a comprehensive technical reference for the Voice Outbound Agent "
        "platform. It covers the system architecture, data flow diagrams, module-level "
        "implementation details, API specifications, deployment procedures, and compliance "
        "requirements. It serves as the authoritative reference for engineers deploying, "
        "maintaining, or extending the system and for operations teams responsible for day-to-day "
        "campaign management."
    )

    h2(doc, "1.2  Scope and Audience")
    para(doc,
        "This document covers all eleven implemented modules, from the database schema through "
        "to the evaluation and load-test suite. It assumes familiarity with Python, FastAPI, "
        "PostgreSQL, Docker, and REST/WebSocket APIs."
    )
    tbl(doc,
        ["Audience", "Relevant Sections"],
        [
            ["Backend engineers", "All sections"],
            ["DevOps / Infrastructure", "Sections 4, 7, Appendix A"],
            ["Product owners / Operations", "Sections 2, 3, 8"],
            ["QA / Test engineers", "Sections 5, 6, and Section 5.11 (evals)"],
        ]
    )

    h2(doc, "1.3  Document Conventions")
    bul(doc, "Monospace text denotes code, file paths, environment variable names, or command-line instructions.")
    bul(doc, "Bold text within prose highlights key terms on first use.")
    bul(doc, "Light-grey shaded boxes contain source code excerpts or command sequences.")
    bul(doc, "Diagram boxes use box-drawing characters to represent system components and data flows.")
    bul(doc, "Figure captions appear in italics below each diagram.")

    h2(doc, "1.4  Related Documents")
    bul(doc, "PRD.md — Product Requirements Document (project root)")
    bul(doc, "CLAUDE.md — Project AI assistant rules and phase status")
    bul(doc, ".env.example — Annotated environment variable template")
    bul(doc, "alembic/versions/ — Database migration history")

    pb(doc)

    # ── Section 2: System Architecture ─────────────────────────────────────
    h1(doc, "2.  System Architecture")

    h2(doc, "2.1  High-Level Overview")
    para(doc,
        "The Voice Outbound Agent is a production-grade autonomous system that dials leads from "
        "a PostgreSQL database, conducts natural AI-driven phone conversations, and extracts "
        "structured data from each call for downstream CRM and automation workflows. The system "
        "is built around the Retell AI orchestration platform, which handles the real-time voice "
        "pipeline — speech-to-text transcription, language model inference, and text-to-speech "
        "synthesis — and delivers calls over the Telnyx SIP carrier network."
    )
    para(doc,
        "The backend is implemented in Python 3.13 using FastAPI for the webhook receiver and "
        "dashboard API, and Redis with RQ for background job processing. Post-call analysis is "
        "performed by Claude Sonnet via the Anthropic API, producing structured JSONB results "
        "written directly to PostgreSQL. Campaign management and live call monitoring are provided "
        "through a Next.js 14 dashboard with real-time WebSocket updates."
    )

    diag(doc,
        " ┌──────────────────────────────────────────────────────────────────────┐\n"
        " │  PostgreSQL 16  —  agent_operations schema                           │\n"
        " │  campaigns · leads · dnc_registry · call_logs · call_transcripts     │\n"
        " └───────────────────────────┬──────────────────────────────────────────┘\n"
        "                             │  Batch query  (DNC-filtered · timezone-gated)\n"
        "                             ▼\n"
        " ┌──────────────────────────────────────────────────────────────────────┐\n"
        " │  Dialing Worker  (app/dialing_worker/)                               │\n"
        " │  Redis / RQ  ·  1 Call Per Second  ·  08:00 – 21:00 local time       │\n"
        " └───────────────────────────┬──────────────────────────────────────────┘\n"
        "                             │  POST /v2/create-call\n"
        "                             ▼\n"
        " ┌──────────────────────────────────────────────────────────────────────┐\n"
        " │  Retell AI Orchestration  (managed service)                          │\n"
        " │  Silero VAD  ►  Deepgram Nova-3 STT  ►  GPT-4o Realtime LLM         │\n"
        " │                          ►  ElevenLabs v3 TTS  ►  Telnyx SIP        │\n"
        " └────────────┬─────────────────────────────────────────────────────────┘\n"
        "              │  Webhooks  (call_started · call_ended · call_analyzed)\n"
        "              ▼\n"
        " ┌──────────────────────────────────────────────────────────────────────┐\n"
        " │  FastAPI Webhook Receiver  (port 8000)                               │\n"
        " │  HMAC-SHA256 verification  ·  Redis replay protection                │\n"
        " └────────────┬─────────────────────────────────────────────────────────┘\n"
        "              │  RQ job enqueue  (call_analyzed)\n"
        "              ▼\n"
        " ┌──────────────────────────────────────────────────────────────────────┐\n"
        " │  Post-Call Analysis Worker  (app/post_call_analysis/)                │\n"
        " │  Claude Sonnet  ·  Structured JSONB extraction  ·  DNC detection     │\n"
        " └────────────┬─────────────────────────────────────────────────────────┘\n"
        "              │\n"
        "     ┌────────┴──────────┐\n"
        "     ▼                   ▼\n"
        " ┌──────────┐   ┌─────────────────────────────────────────────────────┐\n"
        " │  n8n     │   │  Next.js 14 Dashboard  (port 3000)                  │\n"
        " │  Auto-   │   │  Campaign manager · Live call monitor · KPI charts  │\n"
        " │  mation  │   │  REST API + WebSocket  (port 8000)                  │\n"
        " └──────────┘   └─────────────────────────────────────────────────────┘",
        "Figure 1 — High-Level System Architecture"
    )

    h2(doc, "2.2  Technology Stack")
    para(doc,
        "The platform combines managed AI services — used where latency and quality requirements "
        "are non-negotiable — with open-source infrastructure components to minimise ongoing "
        "licensing costs."
    )
    tbl(doc,
        ["Layer", "Component", "Rationale"],
        [
            ["Orchestration", "Retell AI (managed)", "Best-in-class barge-in detection and turn-taking; ~600ms end-to-end latency"],
            ["Realtime LLM", "GPT-4o Realtime (OpenAI)", "Native multimodal audio input; sub-1ms token generation"],
            ["Post-call LLM", "Claude Sonnet (Anthropic)", "200K context window; reliable structured JSON output; cost-efficient for batch"],
            ["Text-to-Speech", "ElevenLabs v3", "~75ms first-audio latency; <20% AI detection rate on first 5 seconds"],
            ["Speech-to-Text", "Deepgram Nova-3", "Streaming partial transcripts via Retell AI integration"],
            ["Voice Activity Detection", "Silero VAD (open-source, MIT)", "200ms speech onset / 800ms silence offset; deterministic state machine"],
            ["Telephony / SIP", "Telnyx", "Direct carrier interconnect; private edge network; lowest call setup latency"],
            ["Database", "PostgreSQL 16", "JSONB for structured extraction; scoped schema for LLM role isolation"],
            ["API Framework", "FastAPI (Python 3.13)", "Async-native; automatic OpenAPI docs; dependency injection via Depends()"],
            ["Job Queue", "Redis + RQ (open-source)", "Strict 1 CPS rate enforcement; exponential backoff retry built-in"],
            ["Automation", "n8n self-hosted (open-source)", "Post-call CRM push, SMS follow-up, calendar booking via webhook trigger"],
            ["Dashboard", "Next.js 14 (open-source)", "App Router; WebSocket live updates; server-side rendering"],
            ["Containerisation", "Docker Compose", "Single-command full-stack startup; health-check-gated service ordering"],
        ]
    )

    h2(doc, "2.3  Service Topology")
    para(doc,
        "In production, the system runs as six Docker Compose services. In local development, "
        "PostgreSQL and Redis run in containers while the API, worker, and dashboard run as "
        "local processes to allow hot-reload during development."
    )
    diag(doc,
        "  ┌─────────────────────────────────────────────────────────────────────┐\n"
        "  │  Docker Compose Network                                             │\n"
        "  │                                                                     │\n"
        "  │  ┌───────────────┐    ┌───────────────┐    ┌──────────────────┐    │\n"
        "  │  │  postgres:16  │    │  redis:7      │    │  n8n  (latest)   │    │\n"
        "  │  │  port 5432    │    │  port 6379    │    │  port 5678       │    │\n"
        "  │  └───────┬───────┘    └───────┬───────┘    └──────────────────┘    │\n"
        "  │          │ (healthy)           │ (healthy)                          │\n"
        "  │          └──────────┬──────────┘                                   │\n"
        "  │                     │  depends_on (service_healthy)                │\n"
        "  │          ┌──────────┴─────────────────┐                            │\n"
        "  │          ▼                             ▼                            │\n"
        "  │  ┌──────────────────┐    ┌────────────────────┐                    │\n"
        "  │  │  api             │    │  worker            │                    │\n"
        "  │  │  Dockerfile.api  │    │  Dockerfile.worker │                    │\n"
        "  │  │  port 8000       │    │  (no exposed port) │                    │\n"
        "  │  └────────┬─────────┘    └────────────────────┘                    │\n"
        "  │           │ (healthy)                                               │\n"
        "  │           ▼                                                         │\n"
        "  │  ┌──────────────────┐                                               │\n"
        "  │  │  dashboard       │                                               │\n"
        "  │  │  Next.js 14      │                                               │\n"
        "  │  │  port 3000       │                                               │\n"
        "  │  └──────────────────┘                                               │\n"
        "  └─────────────────────────────────────────────────────────────────────┘",
        "Figure 2 — Docker Compose Service Topology"
    )

    h2(doc, "2.4  Database Schema")
    para(doc,
        "All application tables reside in the agent_operations PostgreSQL schema, isolated from "
        "the public schema. This ensures LLM-facing database roles cannot access system tables "
        "or data outside the application boundary. All primary keys are UUID v4 values generated "
        "by the database."
    )
    tbl(doc,
        ["Table", "Primary Key", "Purpose"],
        [
            ["campaigns", "campaign_id  UUID", "Campaign configuration: prompt template, LLM settings, status lifecycle"],
            ["leads", "lead_id  UUID", "Dialing queue: phone number, timezone, campaign FK, status, retry metadata"],
            ["dnc_registry", "dnc_id  UUID", "Do-Not-Call numbers: enforced via SQL NOT EXISTS at every query"],
            ["call_logs", "call_id  UUID", "Immutable call telemetry: start/end times, duration, recording URL"],
            ["call_transcripts", "transcript_id  UUID", "Raw transcript and Claude Sonnet structured extraction (JSONB)"],
        ]
    )
    diag(doc,
        "  campaigns                          leads\n"
        "  ─────────────────────────          ────────────────────────────────\n"
        "  campaign_id   UUID  PK             lead_id        UUID  PK\n"
        "  name          VARCHAR              phone_number   VARCHAR(20)  E.164\n"
        "  status        VARCHAR              first_name     VARCHAR\n"
        "  prompt_template  JSONB             last_name      VARCHAR\n"
        "  llm_config       JSONB             company        VARCHAR\n"
        "  created_at    TIMESTAMPTZ          timezone       VARCHAR  (IANA)\n"
        "                                     campaign_id    FK ──────────────► campaigns\n"
        "                                     status         VARCHAR\n"
        "                                     retry_count    INTEGER\n"
        "                                     next_retry_at  TIMESTAMPTZ\n"
        "                                     custom_vars    JSONB\n"
        "                                     created_at     TIMESTAMPTZ\n"
        "\n"
        "  dnc_registry                            │\n"
        "  ─────────────────────────               │ (NOT EXISTS check at query time)\n"
        "  dnc_id        UUID  PK                  │\n"
        "  phone_number  VARCHAR  UNIQUE            ▼\n"
        "  added_at      TIMESTAMPTZ          call_logs\n"
        "  source        VARCHAR              ────────────────────────────────\n"
        "                                     call_id        UUID  PK\n"
        "                                     lead_id        FK ──────────────► leads\n"
        "                                     retell_call_id VARCHAR\n"
        "                                     start_time     TIMESTAMPTZ\n"
        "                                     end_time       TIMESTAMPTZ\n"
        "                                     duration_sec   INTEGER\n"
        "                                     disconnect_reason  VARCHAR\n"
        "                                     recording_url  VARCHAR\n"
        "                                     created_at     TIMESTAMPTZ\n"
        "                                          │\n"
        "                                          ▼\n"
        "                                     call_transcripts\n"
        "                                     ────────────────────────────────\n"
        "                                     transcript_id  UUID  PK\n"
        "                                     call_id        FK ──────────────► call_logs\n"
        "                                     raw_transcript TEXT\n"
        "                                     structured_data  JSONB\n"
        "                                     sentiment      VARCHAR\n"
        "                                     created_at     TIMESTAMPTZ",
        "Figure 3 — Database Entity Relationship Diagram"
    )

    h2(doc, "2.5  Security Architecture")
    h3(doc, "Database Role Isolation")
    para(doc,
        "The memories_role PostgreSQL role is the sole role used by application code for database "
        "access. It is granted SELECT, INSERT, and UPDATE on application tables only — no DELETE, "
        "DROP, TRUNCATE, or CREATE privileges are assigned. The dnc_registry table is read-only "
        "for this role; inserts are made by the post-call analysis worker through an elevated "
        "connection using the application owner role during the DNC auto-insert path."
    )
    h3(doc, "Webhook Signature Verification")
    para(doc,
        "Every incoming webhook request from Retell AI is verified with HMAC-SHA256 before any "
        "processing occurs. The x-retell-signature header contains the hex digest of the raw "
        "request body signed with RETELL_WEBHOOK_SECRET. Requests with an invalid signature "
        "receive HTTP 403; requests missing the header entirely receive HTTP 422 from FastAPI's "
        "header validation layer."
    )
    h3(doc, "Replay Protection")
    para(doc,
        "Each processed webhook event is recorded in Redis under the key pattern "
        "webhook:seen:{call_id}:{event_type} with a 600-second TTL. Duplicate deliveries from "
        "Retell's built-in retry mechanism are detected and silently dropped without re-processing. "
        "If Redis is unavailable, the system degrades gracefully: events are processed normally "
        "and a WARNING is logged."
    )

    pb(doc)

    # ── Section 3: Data Flow Diagrams ───────────────────────────────────────
    h1(doc, "3.  Data Flow Diagrams")

    h2(doc, "3.1  Outbound Call Initiation Flow")
    para(doc,
        "The following diagram shows the complete sequence from campaign activation through to a "
        "successfully initiated outbound call, including all compliance gate checks and the "
        "rate-limit enforcement step."
    )
    diag(doc,
        "  Campaign status set to 'active'\n"
        "          │\n"
        "          ▼\n"
        "  ┌────────────────────────────────────────────────────────┐\n"
        "  │  Dialing Worker  (polls every 5 seconds)               │\n"
        "  │                                                        │\n"
        "  │  SELECT l.*                                            │\n"
        "  │  FROM   agent_operations.leads l                       │\n"
        "  │  JOIN   agent_operations.campaigns c                   │\n"
        "  │           ON l.campaign_id = c.campaign_id             │\n"
        "  │  WHERE  l.status = 'pending'                           │\n"
        "  │    AND  l.next_retry_at <= NOW()                       │\n"
        "  │    AND  NOT EXISTS (                                   │\n"
        "  │           SELECT 1                                     │\n"
        "  │           FROM   agent_operations.dnc_registry d       │\n"
        "  │           WHERE  d.phone_number = l.phone_number       │\n"
        "  │         )                                              │\n"
        "  │  ORDER BY l.created_at  LIMIT 50                       │\n"
        "  └────────────────────────────────────────────────────────┘\n"
        "          │\n"
        "          │  For each lead in result set:\n"
        "          ▼\n"
        "  ┌──────────────────────────────────────────────┐\n"
        "  │  Timezone Gate                               │\n"
        "  │  is_within_calling_hours(lead.timezone)      │\n"
        "  │  Allowed window: 08:00 – 21:00 local time    │\n"
        "  └──────────────┬───────────────────────────────┘\n"
        "                 │ YES                 NO\n"
        "                 │                     └──► Skip lead (re-evaluated next poll)\n"
        "                 ▼\n"
        "  ┌──────────────────────────────────────────────┐\n"
        "  │  Retell AI — Create Outbound Call            │\n"
        "  │  POST https://api.retellai.com/v2/create-call│\n"
        "  │  {                                           │\n"
        "  │    from_number: TELNYX_FROM_NUMBER,          │\n"
        "  │    to_number:   lead.phone_number,           │\n"
        "  │    agent_id:    campaign.retell_agent_id,    │\n"
        "  │    retell_llm_dynamic_variables: {           │\n"
        "  │      first_name, company, ...custom_vars     │\n"
        "  │    },                                        │\n"
        "  │    metadata: { lead_id: <uuid> }             │\n"
        "  │  }                                           │\n"
        "  └──────────────┬───────────────────────────────┘\n"
        "                 │\n"
        "                 ▼\n"
        "  UPDATE leads SET status = 'calling'\n"
        "  await asyncio.sleep(1.0)    ←  1 CPS hard gate\n"
        "                 │\n"
        "                 └──► Continue to next lead in batch",
        "Figure 4 — Outbound Call Initiation Data Flow"
    )

    h2(doc, "3.2  Webhook Event Processing Flow")
    para(doc,
        "Retell AI fires HTTP POST requests to the webhook receiver for each call lifecycle event. "
        "The pipeline verifies authenticity before processing and uses Redis to prevent duplicate "
        "database writes from Retell's automatic retry mechanism."
    )
    diag(doc,
        "  Retell AI  ──►  POST /webhook\n"
        "                        │\n"
        "                        ▼\n"
        "          ┌──────────────────────────────────────────┐\n"
        "          │  HMAC-SHA256 Signature Verification      │\n"
        "          │  x-retell-signature header               │\n"
        "          │  key = RETELL_WEBHOOK_SECRET             │\n"
        "          └─────────────┬────────────────────────────┘\n"
        "                        │\n"
        "           Invalid ◄────┤────► Valid\n"
        "              │                    │\n"
        "         403 Forbidden             ▼\n"
        "                       ┌──────────────────────────────────────────┐\n"
        "                       │  Replay Protection  (Redis)              │\n"
        "                       │  key: webhook:seen:{call_id}:{event}     │\n"
        "                       │  TTL: 600 seconds                        │\n"
        "                       └─────────────┬────────────────────────────┘\n"
        "                                     │\n"
        "                      Duplicate ◄────┤────► New event\n"
        "                         │                       │\n"
        "                    200 OK (skip)                ▼\n"
        "                                    ┌───────────────────────────┐\n"
        "                                    │  Event Router             │\n"
        "                                    └─────┬──────┬────┬─────────┘\n"
        "                                          │      │    │\n"
        "                                   call_  │    call_  │  transcript_\n"
        "                                   started│    ended  │   updated\n"
        "                                          │      │    │\n"
        "                                          ▼      ▼    └──► Log only\n"
        "                                   Upsert     Update\n"
        "                                   call_logs  call_logs\n"
        "                                   leads →    leads →\n"
        "                                   'calling'  'completed'\n"
        "                                              or 'failed'\n"
        "\n"
        "                                   call_analyzed\n"
        "                                          │\n"
        "                                          ▼\n"
        "                                   Insert call_transcripts.raw_transcript\n"
        "                                   Enqueue RQ job: analyze_call(call_id)",
        "Figure 5 — Webhook Event Processing Flow"
    )

    h2(doc, "3.3  Post-Call Analysis Pipeline")
    para(doc,
        "After each call, Retell AI generates a full transcript and fires a call_analyzed webhook "
        "event. The post-call analysis worker processes this asynchronously via an RQ job, calling "
        "Claude Sonnet to extract structured business intelligence."
    )
    diag(doc,
        "  RQ Queue  ──►  analyze_call(call_id)\n"
        "                        │\n"
        "                        ▼\n"
        "          ┌──────────────────────────────────────────────┐\n"
        "          │  Load records from PostgreSQL                │\n"
        "          │  · call_transcripts  (raw_transcript text)   │\n"
        "          │  · call_logs         (duration, metadata)    │\n"
        "          │  · leads             (context variables)     │\n"
        "          └──────────────────┬───────────────────────────┘\n"
        "                             │\n"
        "                             ▼\n"
        "          ┌──────────────────────────────────────────────┐\n"
        "          │  Claude Sonnet API Call                      │\n"
        "          │  Model: claude-sonnet-4-6                    │\n"
        "          │  Tool:  extract_call_data (JSON schema)      │\n"
        "          │  Input: raw_transcript text                  │\n"
        "          └──────────────────┬───────────────────────────┘\n"
        "                             │  Returns ExtractionResult\n"
        "                             ▼\n"
        "          ┌──────────────────────────────────────────────┐\n"
        "          │  Two-Layer DNC Detection                     │\n"
        "          │  Layer 1: result.dnc_requested == True       │\n"
        "          │  Layer 2: dnc_keywords.scan(transcript)      │\n"
        "          └──────────────────┬───────────────────────────┘\n"
        "                 No DNC      │      DNC detected\n"
        "                             │             │\n"
        "                             │             ▼\n"
        "                             │  INSERT dnc_registry\n"
        "                             │  UPDATE leads.status = 'failed_dnc'\n"
        "                             │\n"
        "                             ▼\n"
        "          ┌──────────────────────────────────────────────┐\n"
        "          │  Write to PostgreSQL                         │\n"
        "          │  call_transcripts.structured_data  JSONB     │\n"
        "          │  call_transcripts.sentiment        VARCHAR   │\n"
        "          └──────────────────┬───────────────────────────┘\n"
        "                             │\n"
        "                             ▼\n"
        "          POST n8n  N8N_WEBHOOK_URL  (fire-and-forget)\n"
        "          Payload: full ExtractionResult JSON",
        "Figure 6 — Post-Call Analysis Pipeline"
    )

    h2(doc, "3.4  DNC Enforcement Flow")
    para(doc,
        "DNC enforcement operates at two independent points in the system lifecycle, ensuring "
        "that no registered number can be dialled regardless of application state or race "
        "conditions."
    )
    diag(doc,
        "  ┌──────────────────────────────┐   ┌───────────────────────────────────────┐\n"
        "  │  BEFORE DIAL                 │   │  AFTER CALL                           │\n"
        "  │  SQL NOT EXISTS gate         │   │  Post-call analysis detection         │\n"
        "  └──────────────┬───────────────┘   └──────────────────┬────────────────────┘\n"
        "                 │                                       │\n"
        "  SELECT ... FROM leads              Claude Sonnet flag: dnc_requested = True\n"
        "  WHERE NOT EXISTS (                               OR\n"
        "    SELECT 1 FROM dnc_registry       Keyword scanner:  'remove me from list'\n"
        "    WHERE phone_number = l.phone                        'stop calling me'\n"
        "  )                                                     'do not call'\n"
        "                 │                                       │\n"
        "     Number in DNC registry                             │\n"
        "                 │                                       ▼\n"
        "                 ▼                   INSERT dnc_registry (source='caller_request')\n"
        "  Lead excluded from query result    UPDATE leads.status = 'failed_dnc'\n"
        "  Never reaches the worker           Number blocked from ALL future campaigns",
        "Figure 7 — DNC Enforcement at Both Lifecycle Stages"
    )

    pb(doc)

    # ── Section 4: Local Development Setup ──────────────────────────────────
    h1(doc, "4.  Local Development Setup")

    h2(doc, "4.1  Prerequisites")
    tbl(doc,
        ["Tool", "Minimum Version", "Notes"],
        [
            ["Python", "3.13", "Use pyenv or the official installer at python.org"],
            ["Node.js", "20 LTS", "Required for the Next.js dashboard only"],
            ["Docker Desktop", "4.x", "Runs PostgreSQL, Redis, and n8n containers"],
            ["Git", "2.x", "Standard version control"],
        ]
    )

    h2(doc, "4.2  Repository Setup")
    code(doc,
        "git clone https://github.com/abhi-30702/voice-outbound-agent.git\n"
        "cd voice-outbound-agent\n"
        "\n"
        "# Create and activate the Python virtual environment\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate          # Windows PowerShell\n"
        "# source .venv/bin/activate       # macOS / Linux\n"
        "\n"
        "# Install all Python dependencies\n"
        "pip install -r requirements.txt"
    )

    h2(doc, "4.3  Environment Configuration")
    para(doc,
        "Copy the annotated example file and populate it with real API credentials. The system "
        "will not make any external API calls until real keys are provided."
    )
    code(doc,
        "copy .env.example .env    # Windows\n"
        "# cp .env.example .env   # macOS / Linux"
    )
    para(doc, "Minimum set of variables required to run locally:")
    tbl(doc,
        ["Variable", "Required", "Description"],
        [
            ["DATABASE_URL", "Yes", "postgresql+asyncpg://postgres:password@localhost:5432/voice_agent"],
            ["REDIS_URL", "No", "Default: redis://localhost:6379/0"],
            ["RETELL_API_KEY", "Yes", "From retellai.com → Dashboard → API Keys"],
            ["RETELL_WEBHOOK_SECRET", "Yes", "From retellai.com → Webhook Settings"],
            ["TELNYX_API_KEY", "Yes", "From telnyx.com → Auth → API Keys"],
            ["TELNYX_FROM_NUMBER", "Yes", "Purchased Telnyx number in E.164 format"],
            ["ELEVENLABS_API_KEY", "Yes", "From elevenlabs.io → Profile → API Key"],
            ["ELEVENLABS_VOICE_ID", "Yes", "Voice ID from ElevenLabs voice library"],
            ["OPENAI_API_KEY", "Yes", "From platform.openai.com → API Keys"],
            ["ANTHROPIC_API_KEY", "Yes", "From console.anthropic.com → API Keys"],
            ["WEBHOOK_BASE_URL", "Yes", "Publicly reachable HTTPS URL; use ngrok for local development"],
        ]
    )

    h2(doc, "4.4  Starting the Full Stack")
    para(doc, "Step 1 — Start infrastructure containers:")
    code(doc, "docker compose up -d postgres redis")

    para(doc, "Step 2 — Apply database migrations:")
    code(doc, "alembic upgrade head")

    para(doc, "Step 3 — Start the API server (Terminal 1):")
    code(doc,
        ".venv\\Scripts\\python -m uvicorn app.webhook_receiver.main:app \\\n"
        "    --host 0.0.0.0 --port 8000 --reload"
    )

    para(doc, "Step 4 — Start the background worker (Terminal 2):")
    code(doc, ".venv\\Scripts\\python scripts\\run_worker.py")

    para(doc, "Step 5 — Start the Next.js dashboard (Terminal 3):")
    code(doc,
        "cd app\\dashboard\n"
        "npm install\n"
        "npm run dev"
    )

    para(doc, "Step 6 — (Optional) Start n8n automation:")
    code(doc, "docker compose up -d n8n")

    h2(doc, "4.5  Verification and Health Checks")
    tbl(doc,
        ["Service", "URL", "Expected"],
        [
            ["API — Swagger docs", "http://localhost:8000/docs", "OpenAPI UI with all endpoints listed"],
            ["API — DB health probe", "http://localhost:8000/health/db", "{\"status\": \"ready\", \"migration_version\": \"001\"}"],
            ["Next.js dashboard", "http://localhost:3000", "Campaign management page renders"],
            ["n8n editor", "http://localhost:5678", "Login page (default: admin / changeme)"],
        ]
    )

    h2(doc, "4.6  Running the Test Suite")
    para(doc, "All mock and unit tests — no live database or external API calls required:")
    code(doc, "pytest tests/ --ignore=tests/evals/test_dnc_regression.py -v")

    para(doc, "Evaluation suite only:")
    code(doc, "pytest tests/evals/ --ignore=tests/evals/test_dnc_regression.py -v")

    para(doc, "DNC regression test — requires live PostgreSQL with seeded test data:")
    code(doc, "pytest tests/evals/test_dnc_regression.py -v")

    para(doc, "Standalone KPI check against the live database:")
    code(doc, "python tests/evals/kpi_check.py")

    pb(doc)

    # ── Section 5: Module Reference ──────────────────────────────────────────
    h1(doc, "5.  Module Reference")
    para(doc,
        "The system is organised into eleven modules, each scoped to a single directory. Each "
        "module has a clearly defined responsibility and communicates with other modules only "
        "through well-defined interfaces — database tables, Redis queues, or HTTP endpoints. "
        "This isolation ensures that modules can be developed, tested, and deployed independently."
    )

    # 5.1 db-schema
    h2(doc, "5.1  db-schema  (app/db/ + alembic/)")
    para(doc,
        "Provides the PostgreSQL schema definition, Alembic migration history, SQLAlchemy ORM "
        "models, async session management, parameterized query helpers, and database health probes. "
        "This module defines the data contract for the entire system."
    )
    h3(doc, "Key Files")
    tbl(doc,
        ["File", "Responsibility"],
        [
            ["app/db/session.py", "Async SQLAlchemy engine singleton and session factory"],
            ["app/db/dependencies.py", "FastAPI Depends(get_db) dependency injection"],
            ["app/db/queries.py", "Parameterized helpers: DNC check, pending lead fetch, call lookup"],
            ["app/db/init_db.py", "Schema creation, connectivity health check, migration status probe"],
            ["app/models/campaign.py", "Campaign ORM model with JSONB prompt_template and llm_config"],
            ["app/models/contact.py", "Lead (Contact) ORM model with ContactStatus enum and retry fields"],
            ["app/models/call.py", "Call log ORM model — effectively INSERT-only after creation"],
            ["app/models/transcript.py", "Call transcript model with structured_data JSONB column"],
            ["app/models/dnc_entry.py", "DNC registry model — SELECT and INSERT only via memories_role"],
            ["alembic/versions/001_*.py", "Initial migration: creates agent_operations schema and all 5 tables"],
        ]
    )
    h3(doc, "Code Snippet — DNC-Filtered Lead Fetch")
    code(doc,
        "# app/db/queries.py\n"
        "async def get_pending_leads(\n"
        "    session: AsyncSession, limit: int = 50\n"
        ") -> list[Contact]:\n"
        "    result = await session.execute(\n"
        "        select(Contact)\n"
        "        .where(\n"
        "            Contact.status == ContactStatus.PENDING,\n"
        "            Contact.next_retry_at <= func.now(),\n"
        "            ~exists(\n"
        "                select(DncEntry.id).where(\n"
        "                    DncEntry.phone_number == Contact.phone_number\n"
        "                )\n"
        "            ),\n"
        "        )\n"
        "        .order_by(Contact.created_at)\n"
        "        .limit(limit)\n"
        "    )\n"
        "    return list(result.scalars().all())"
    )
    h3(doc, "Security Notes")
    bul(doc, "All tables are in the agent_operations schema — not public.")
    bul(doc, "The memories_role role has SELECT/INSERT/UPDATE only — no DELETE, DROP, or TRUNCATE.")
    bul(doc, "All queries use SQLAlchemy parameterization — string interpolation in SQL is prohibited.")

    # 5.2 dialing-worker
    h2(doc, "5.2  dialing-worker  (app/dialing_worker/)")
    para(doc,
        "The autonomous outbound dialing worker. Polls PostgreSQL every five seconds for pending "
        "leads that have passed the DNC SQL gate and timezone gate, dispatches calls via the "
        "Retell AI API, and enforces a strict one-call-per-second rate limit."
    )
    h3(doc, "Key Files")
    tbl(doc,
        ["File", "Responsibility"],
        [
            ["app/dialing_worker/worker.py", "DialerWorker class: dial_batch() and _dispatch_call()"],
            ["app/dialing_worker/timezone_utils.py", "is_within_calling_hours(): IANA timezone to local time comparison"],
            ["scripts/run_worker.py", "Entry point: initialises database engine and starts the worker loop"],
        ]
    )
    h3(doc, "Code Snippet — Rate-Limited Dispatch Loop")
    code(doc,
        "# app/dialing_worker/worker.py\n"
        "async def dial_batch(self, session: AsyncSession) -> int:\n"
        "    leads = await get_pending_leads(session, limit=50)\n"
        "    dialable = [\n"
        "        lead for lead in leads\n"
        "        if is_within_calling_hours(lead.timezone)\n"
        "    ]\n"
        "    dispatched = 0\n"
        "    for lead in dialable:\n"
        "        await self._dispatch_call(session, lead)\n"
        "        await asyncio.sleep(1.0)   # strict 1 CPS enforcement\n"
        "        dispatched += 1\n"
        "    return dispatched"
    )
    h3(doc, "Operational Notes")
    bul(doc, "Batch size: 50 leads per poll cycle.")
    bul(doc, "Poll interval: 5 seconds between cycles.")
    bul(doc, "Rate limit: exactly 1 call per second via asyncio.sleep(1.0) after each dispatch.")
    bul(doc, "Retry backoff: 60s → 120s → 300s; maximum 3 retries before status is set to 'failed'.")
    bul(doc, "Run with: .venv\\Scripts\\python scripts\\run_worker.py")

    # 5.3 retell-integration
    h2(doc, "5.3  retell-integration  (app/retell_integration/)")
    para(doc,
        "Provides the Retell AI API client used by the dialing worker to create outbound calls "
        "and manage agent configurations. The sync_agent() function is idempotent — it creates "
        "a new Retell agent if one does not exist for the campaign, or updates it if the prompt "
        "template or LLM config has changed."
    )
    h3(doc, "Code Snippet — Creating an Outbound Call")
    code(doc,
        "# app/retell_integration/client.py\n"
        "async def create_call(\n"
        "    self, to_number: str, agent_id: str,\n"
        "    dynamic_variables: dict, lead_id: str\n"
        ") -> dict:\n"
        "    payload = {\n"
        "        \"from_number\": settings.TELNYX_FROM_NUMBER,\n"
        "        \"to_number\": to_number,\n"
        "        \"agent_id\": agent_id,\n"
        "        \"retell_llm_dynamic_variables\": dynamic_variables,\n"
        "        \"metadata\": {\"lead_id\": lead_id},\n"
        "    }\n"
        "    async with httpx.AsyncClient() as client:\n"
        "        resp = await client.post(\n"
        "            \"https://api.retellai.com/v2/create-call\",\n"
        "            json=payload,\n"
        "            headers={\"Authorization\": f\"Bearer {self.api_key}\"},\n"
        "        )\n"
        "        resp.raise_for_status()\n"
        "        return resp.json()"
    )
    para(doc,
        "Before running any campaign, call sync_agent() to register or update the Retell agent:"
    )
    code(doc,
        "from app.retell_integration.agent_manager import sync_agent\n"
        "import asyncio\n"
        "asyncio.run(sync_agent(campaign_id=\"your-campaign-uuid\"))"
    )

    # 5.4 webhook-receiver
    h2(doc, "5.4  webhook-receiver  (app/webhook_receiver/)")
    para(doc,
        "A FastAPI application that receives all Retell AI lifecycle events. Every request is "
        "verified with HMAC-SHA256 before processing. The service is stateless — it reads and "
        "writes only to PostgreSQL and Redis — making horizontal scaling straightforward."
    )
    h3(doc, "Key Files")
    tbl(doc,
        ["File", "Responsibility"],
        [
            ["app/webhook_receiver/main.py", "FastAPI app entry point and lifespan handler"],
            ["app/webhook_receiver/router.py", "POST /webhook endpoint handler"],
            ["app/webhook_receiver/services/signature_verifier.py", "HMAC-SHA256 verification"],
            ["app/webhook_receiver/services/event_dispatcher.py", "Routes events to per-event handlers"],
            ["app/webhook_receiver/services/queue_service.py", "Enqueues RQ jobs for post-call analysis"],
        ]
    )
    h3(doc, "Code Snippet — HMAC Signature Verification")
    code(doc,
        "# app/webhook_receiver/services/signature_verifier.py\n"
        "import hashlib, hmac\n"
        "\n"
        "def verify_signature(\n"
        "    body: bytes, signature: str, secret: str\n"
        ") -> bool:\n"
        "    expected = hmac.new(\n"
        "        key=secret.encode(\"utf-8\"),\n"
        "        msg=body,\n"
        "        digestmod=hashlib.sha256,\n"
        "    ).hexdigest()\n"
        "    return hmac.compare_digest(expected, signature)"
    )
    h3(doc, "Event → Action Mapping")
    tbl(doc,
        ["Event", "Action"],
        [
            ["call_started", "Upsert call_logs row; set leads.status = 'calling'"],
            ["call_ended", "Update call_logs with duration, disconnect_reason, recording_url; set leads status"],
            ["call_analyzed", "Insert raw_transcript; enqueue analyze_call(call_id) RQ job"],
            ["transcript_updated", "Broadcast to WebSocket clients only — no database write"],
        ]
    )

    # 5.5 post-call-analysis
    h2(doc, "5.5  post-call-analysis  (app/post_call_analysis/)")
    para(doc,
        "An RQ background worker that processes call transcripts using the Anthropic API. It "
        "extracts structured business intelligence and writes JSONB results to call_transcripts. "
        "It also performs two-layer DNC detection and triggers n8n automation."
    )
    h3(doc, "Code Snippet — Claude Sonnet Extraction")
    code(doc,
        "# app/post_call_analysis/worker.py\n"
        "def _call_claude(transcript: str) -> ExtractionResult:\n"
        "    client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)\n"
        "    response = client.messages.create(\n"
        "        model=\"claude-sonnet-4-6\",\n"
        "        max_tokens=1024,\n"
        "        tools=[EXTRACTION_TOOL_SCHEMA],\n"
        "        tool_choice={\"type\": \"tool\", \"name\": \"extract_call_data\"},\n"
        "        messages=[{\"role\": \"user\", \"content\": transcript}],\n"
        "    )\n"
        "    tool_block = next(\n"
        "        b for b in response.content if b.type == \"tool_use\"\n"
        "    )\n"
        "    return ExtractionResult(**tool_block.input)"
    )
    h3(doc, "Retry Policy")
    bul(doc, "3 attempts with exponential backoff: 60s → 120s → 300s.")
    bul(doc, "On exhaustion, writes {\"failed_analysis\": true, \"error\": \"...\"} to structured_data.")

    # 5.6 vad-pipeline
    h2(doc, "5.6  vad-pipeline  (app/vad_pipeline/)")
    para(doc,
        "Wraps the open-source Silero VAD model with a four-state speech detection machine. "
        "The pipeline determines exactly when a person has started and stopped speaking, enabling "
        "the AI to respond at the right moment and handle user interruptions gracefully."
    )
    h3(doc, "VAD State Machine")
    diag(doc,
        "  QUIET ──(audio > threshold, sustained 200ms)──► STARTING\n"
        "                                                       │\n"
        "  STOPPING ◄──(silence > 800ms)── SPEAKING ◄──(200ms confirmed)\n"
        "      │                               │\n"
        "      │ (LLM audio starts)            └──► stream to STT\n"
        "      ▼\n"
        "  QUIET\n"
        "\n"
        "  Interrupt path:\n"
        "  Agent SPEAKING ──(user audio detected)──► HALT TTS stream\n"
        "                                        ──► truncate LLM context to last word\n"
        "                                        ──► STARTING (user turn)",
        "Figure 8 — VAD Four-State Machine"
    )

    # 5.7 conversation-prompts
    h2(doc, "5.7  conversation-prompts  (app/conversation_prompts/)")
    para(doc,
        "Provides the system prompt library with three pre-built persona templates, "
        "an objection handling script library, and a prompt renderer and validator. "
        "All prompts must comply with voice prompt rules — maximum 12 words per sentence, "
        "no bullets or special characters, acronyms spelled phonetically, every step "
        "waits for user acknowledgement before advancing."
    )
    h3(doc, "Prompt Template JSON Structure")
    code(doc,
        "{\n"
        "  \"persona\": {\n"
        "    \"name\": \"Priya\",\n"
        "    \"tone\": \"friendly, professional, unhurried\"\n"
        "  },\n"
        "  \"objective\": \"Confirm the lead's interest in the product demo.\",\n"
        "  \"flow\": [\n"
        "    \"Greeting and identity verification — wait for response\",\n"
        "    \"Purpose statement, one sentence — wait for acknowledgement\",\n"
        "    \"Core qualifying question\",\n"
        "    \"Handle objection or confirm next step\",\n"
        "    \"Polite close\"\n"
        "  ],\n"
        "  \"objections\": {\n"
        "    \"busy\": \"Of course. Can I call at a better time?\",\n"
        "    \"not interested\": \"Absolutely. Have a great day.\",\n"
        "    \"angry\": \"I understand. I'll send an email instead.\"\n"
        "  }\n"
        "}"
    )

    # 5.8 dashboard
    h2(doc, "5.8  dashboard  (app/dashboard/ + app/dashboard_api/)")
    para(doc,
        "A Next.js 14 frontend backed by a FastAPI REST and WebSocket API. Provides real-time "
        "call monitoring with transcript streaming, campaign management, lead CSV upload, "
        "and KPI reporting over configurable time ranges."
    )
    h3(doc, "Dashboard Routes")
    tbl(doc,
        ["Route", "Description"],
        [
            ["/campaigns", "List campaigns; create, activate, or pause a campaign"],
            ["/leads", "Filter leads by campaign and status; upload a CSV file of leads"],
            ["/live", "Real-time call monitor — WebSocket-driven with live transcript feed"],
        ]
    )

    # 5.9 n8n-flows
    h2(doc, "5.9  n8n-flows  (n8n-flows/)")
    para(doc,
        "Post-call automation workflows implemented as n8n JSON exports. The primary workflow "
        "fires on the post-call webhook from the analysis worker and executes a six-node "
        "sequence: parse the result, append to Google Sheets, conditionally send an SMS "
        "follow-up, and optionally book a calendar appointment."
    )
    h3(doc, "Importing the Workflow")
    num(doc, "Navigate to http://localhost:5678 and log in.")
    num(doc, "Go to Workflows → Import from file.")
    num(doc, "Select n8n-flows/post-call-flow.json.")
    num(doc, "Configure Google Sheets, Twilio, and Google Calendar credentials in the Credentials section.")
    num(doc, "Set the webhook trigger URL in the n8n workflow to match N8N_WEBHOOK_URL in .env.")
    num(doc, "Activate the workflow using the toggle in the top-right corner.")

    # 5.10 docker-compose
    h2(doc, "5.10  docker-compose  (project root)")
    para(doc,
        "Full-stack Docker Compose orchestration for all six services. Health checks are defined "
        "for PostgreSQL and Redis; dependent services will not start until their dependencies "
        "report healthy. This prevents the api and worker from attempting database connections "
        "before the database is ready."
    )
    h3(doc, "Code Snippet — Health-Check-Gated Service Startup")
    code(doc,
        "# docker-compose.yml (excerpt)\n"
        "services:\n"
        "  postgres:\n"
        "    image: postgres:16-alpine\n"
        "    healthcheck:\n"
        "      test: [\"CMD-SHELL\", \"pg_isready\"]\n"
        "      interval: 10s\n"
        "      retries: 5\n"
        "\n"
        "  redis:\n"
        "    image: redis:7-alpine\n"
        "    healthcheck:\n"
        "      test: [\"CMD\", \"redis-cli\", \"ping\"]\n"
        "      interval: 10s\n"
        "      retries: 5\n"
        "\n"
        "  api:\n"
        "    build:\n"
        "      context: .\n"
        "      dockerfile: Dockerfile.api\n"
        "    ports: [\"8000:8000\"]\n"
        "    depends_on:\n"
        "      postgres: { condition: service_healthy }\n"
        "      redis:    { condition: service_healthy }"
    )

    # 5.11 evals
    h2(doc, "5.11  evals  (tests/evals/)")
    para(doc,
        "A dedicated evaluation suite covering DNC compliance regression, rate limit "
        "verification, timezone gate correctness, structured output schema validation, "
        "webhook signature security, and a CPS load test. All tests except the DNC "
        "regression test run without a live database."
    )
    h3(doc, "Eval Test Inventory")
    tbl(doc,
        ["Test File", "Coverage", "Live DB"],
        [
            ["test_dnc_regression.py", "100 leads / 10 DNC → asserts 0 DNC numbers are dialled", "Yes"],
            ["test_worker_rate_limit.py", "10 leads: asserts elapsed ≥ 9.0s and create_call called 10×", "No"],
            ["test_timezone_gate.py", "5 timezone scenarios with freezegun: asserts correct allow/block decisions", "No"],
            ["test_structured_output.py", "3 transcript scenarios: validates ExtractionResult Pydantic schema and field values", "No"],
            ["test_signature_verification.py", "Valid HMAC → 200, tampered body → 403, missing header → 422", "No"],
            ["locustfile.py", "CPS load test: 10 leads dispatched, asserts max CPS < 1.1", "No"],
            ["kpi_check.py", "Queries live DB for KPI threshold breaches; exits with code 1 on any breach", "Yes"],
        ]
    )

    pb(doc)

    # ── Section 6: API Reference ─────────────────────────────────────────────
    h1(doc, "6.  API Reference")

    h2(doc, "6.1  Webhook Endpoint")
    para(doc,
        "Base URL is configured via the WEBHOOK_BASE_URL environment variable and must be a "
        "publicly reachable URL that is registered in the Retell AI dashboard."
    )

    h3(doc, "POST /webhook")
    para(doc,
        "The single inbound endpoint for all Retell AI call lifecycle events. Every request "
        "must carry a valid HMAC-SHA256 signature in the x-retell-signature header."
    )
    tbl(doc,
        ["HTTP Status", "Condition"],
        [
            ["200 OK", "Event processed — response body: {\"status\": \"ok\"}"],
            ["403 Forbidden", "HMAC signature is invalid or does not match the body"],
            ["422 Unprocessable Entity", "x-retell-signature header is absent, or request body is malformed"],
            ["500 Internal Server Error", "Unexpected processing error — Retell AI will retry automatically"],
        ]
    )
    h3(doc, "Example Webhook Payloads")
    code(doc,
        "// call_started\n"
        "{\n"
        "  \"event\": \"call_started\",\n"
        "  \"call_id\": \"call_abc123\",\n"
        "  \"agent_id\": \"agent_xyz789\",\n"
        "  \"metadata\": { \"lead_id\": \"550e8400-e29b-41d4-a716-446655440000\" }\n"
        "}\n"
        "\n"
        "// call_ended\n"
        "{\n"
        "  \"event\": \"call_ended\",\n"
        "  \"call_id\": \"call_abc123\",\n"
        "  \"duration_ms\": 145000,\n"
        "  \"disconnect_reason\": \"user_hangup\",\n"
        "  \"recording_url\": \"https://retell.ai/recordings/call_abc123.mp3\"\n"
        "}\n"
        "\n"
        "// call_analyzed\n"
        "{\n"
        "  \"event\": \"call_analyzed\",\n"
        "  \"call_id\": \"call_abc123\",\n"
        "  \"transcript\": \"Agent: Hello, may I speak with...\\nUser: Yes, speaking.\"\n"
        "}"
    )

    h2(doc, "6.2  Dashboard REST API")
    para(doc, "Base URL: http://localhost:8000/api  (configured via NEXT_PUBLIC_API_URL)")
    tbl(doc,
        ["Method", "Path", "Description"],
        [
            ["GET", "/campaigns", "List all campaigns"],
            ["POST", "/campaigns", "Create a campaign (body: CampaignCreate JSON)"],
            ["PATCH", "/campaigns/{id}/status", "Update status to active | paused | completed"],
            ["GET", "/leads", "List leads; filter by campaign_id and/or status query params"],
            ["POST", "/leads/upload", "Upload CSV file of leads (multipart/form-data; optional campaign_id param)"],
            ["POST", "/leads/assign", "Assign a set of leads to a campaign (body: LeadAssign JSON)"],
            ["GET", "/calls/active", "List all calls currently in CALLING status"],
            ["GET", "/kpi", "KPI summary; range query param: today (default) | 7d | 30d"],
        ]
    )
    h3(doc, "Code Snippet — Creating a Campaign")
    code(doc,
        "curl -X POST http://localhost:8000/api/campaigns \\\n"
        "  -H 'Content-Type: application/json' \\\n"
        "  -d '{\n"
        "    \"name\": \"Q3 2026 Outreach\",\n"
        "    \"prompt_template\": {\n"
        "      \"persona\": {\"name\": \"Priya\", \"tone\": \"professional\"},\n"
        "      \"objective\": \"Qualify interest in the product demo.\"\n"
        "    },\n"
        "    \"llm_config\": {\n"
        "      \"retell_agent_id\": \"agent_xyz789\",\n"
        "      \"model\": \"gpt-4o-realtime-preview\"\n"
        "    }\n"
        "  }'"
    )

    h2(doc, "6.3  WebSocket — Live Call Feed")
    para(doc,
        "The dashboard connects to /ws/calls to receive real-time call events. The server "
        "broadcasts a JSON message to all connected clients whenever a webhook event is processed. "
        "Clients should implement automatic reconnection on disconnect."
    )
    code(doc,
        "// Endpoint\n"
        "ws://localhost:8000/ws/calls\n"
        "\n"
        "// Example message broadcast to all connected clients\n"
        "{\n"
        "  \"event_type\": \"call_started\",\n"
        "  \"call_id\": \"call_abc123\",\n"
        "  \"lead_id\": \"550e8400-e29b-41d4-a716-446655440000\",\n"
        "  \"timestamp\": \"2026-05-11T10:30:00Z\"\n"
        "}\n"
        "\n"
        "// JavaScript connection with auto-reconnect\n"
        "function connect() {\n"
        "  const ws = new WebSocket('ws://localhost:8000/ws/calls');\n"
        "  ws.onmessage = (e) => { const d = JSON.parse(e.data); updateUI(d); };\n"
        "  ws.onclose   = () => setTimeout(connect, 2000);\n"
        "}"
    )

    h2(doc, "6.4  Structured Output Schema")
    para(doc,
        "The post-call analysis worker writes the following JSON structure to "
        "call_transcripts.structured_data (JSONB column) after each call:"
    )
    code(doc,
        "{\n"
        "  \"call_outcome\":     \"interested\",\n"
        "  \"callback_time\":    null,\n"
        "  \"objections_raised\": [\"pricing\", \"timing\"],\n"
        "  \"next_action\":      \"Send product brochure and schedule demo\",\n"
        "  \"summary\":          \"Lead expressed strong interest and requested a follow-up demo.\",\n"
        "  \"sentiment\":        \"positive\",\n"
        "  \"sentiment_reason\": \"Caller used affirmative language and asked clarifying questions.\",\n"
        "  \"lead_temperature\": \"hot\",\n"
        "  \"dnc_requested\":    false\n"
        "}"
    )
    tbl(doc,
        ["Field", "Type", "Possible Values"],
        [
            ["call_outcome", "string", "interested | not_interested | callback_requested | dnc_request | no_answer | other"],
            ["callback_time", "string | null", "Verbatim callback time if stated by the lead; null otherwise"],
            ["objections_raised", "list[string]", "Distinct objection types raised during the call"],
            ["next_action", "string", "Recommended follow-up action for the sales team"],
            ["summary", "string", "1–2 sentence plain-language call outcome summary"],
            ["sentiment", "string", "positive | neutral | negative"],
            ["sentiment_reason", "string", "Explanation for the sentiment classification"],
            ["lead_temperature", "string", "hot | warm | cold"],
            ["dnc_requested", "bool", "True if the lead asked to be removed from call lists"],
        ]
    )

    pb(doc)

    # ── Section 7: Deployment Guide ───────────────────────────────────────────
    h1(doc, "7.  Deployment Guide")

    h2(doc, "7.1  Docker Compose Full-Stack Startup")
    para(doc,
        "All six services are defined in docker-compose.yml. Build and start the complete stack "
        "with a single command:"
    )
    code(doc, "docker compose up --build -d")
    para(doc, "Apply database migrations after the first startup (only required once):")
    code(doc, "docker compose run --rm api alembic upgrade head")
    h3(doc, "Common Operational Commands")
    code(doc,
        "docker compose logs -f api          # stream API server logs\n"
        "docker compose logs -f worker       # stream background worker logs\n"
        "docker compose ps                   # display all service health statuses\n"
        "docker compose restart api          # restart API after environment variable changes\n"
        "docker compose down                 # stop all services, preserve volumes\n"
        "docker compose down -v              # stop all services and remove all volumes"
    )

    h2(doc, "7.2  Production Environment Variables")
    tbl(doc,
        ["Variable", "Required Production Value"],
        [
            ["RETELL_API_KEY", "Real key from retellai.com → API Keys"],
            ["RETELL_WEBHOOK_SECRET", "Real secret from retellai.com → Webhook Settings"],
            ["TELNYX_API_KEY", "Real key from telnyx.com → Auth → API Keys"],
            ["TELNYX_FROM_NUMBER", "Purchased Telnyx number in E.164 format (e.g. +12125551234)"],
            ["ELEVENLABS_API_KEY", "Real key from elevenlabs.io → Profile → API Key"],
            ["ELEVENLABS_VOICE_ID", "Voice ID selected from the ElevenLabs voice library"],
            ["OPENAI_API_KEY", "Real key from platform.openai.com → API Keys"],
            ["ANTHROPIC_API_KEY", "Real key from console.anthropic.com → API Keys"],
            ["DATABASE_URL", "postgresql+asyncpg://user:pass@db-host:5432/voice_agent"],
            ["REDIS_URL", "redis://redis-host:6379/0"],
            ["WEBHOOK_BASE_URL", "Publicly reachable HTTPS URL (e.g. https://agent.yourdomain.com)"],
            ["POSTGRES_PASSWORD", "Strong random password — replace default 'password'"],
            ["N8N_BASIC_AUTH_PASSWORD", "Strong random password — replace default 'changeme'"],
        ]
    )

    h2(doc, "7.3  Retell AI Webhook Configuration")
    num(doc, "Log in to retellai.com and navigate to Dashboard → Webhooks.")
    num(doc, "Set the Webhook URL to: https://your-domain.com/webhook")
    num(doc, "Copy the Webhook Secret and set it as RETELL_WEBHOOK_SECRET in .env.")
    num(doc, "Enable all four event types: call_started, call_ended, call_analyzed, transcript_updated.")
    num(doc, "Click Save and send a test event to verify the endpoint returns HTTP 200.")

    h2(doc, "7.4  Going-Live Checklist")
    bul(doc, "All production environment variables populated and verified in .env.")
    bul(doc, "docker compose up --build -d succeeds with all 6 services reporting healthy.")
    bul(doc, "alembic upgrade head completed against the production database.")
    bul(doc, "WEBHOOK_BASE_URL is publicly reachable — test with: curl https://your-domain.com/health/db")
    bul(doc, "Retell AI webhook URL configured; test event returns HTTP 200.")
    bul(doc, "sync_agent() called for each campaign before activation.")
    bul(doc, "At least one campaign set to status='active'.")
    bul(doc, "Leads imported with status='pending' and valid IANA timezone strings.")
    bul(doc, "DNC registry pre-populated with any existing suppression list data.")
    bul(doc, "n8n workflow imported, credentials configured, and workflow activated.")
    bul(doc, "End-to-end test completed: dial one test number; verify call_logs and call_transcripts rows.")

    pb(doc)

    # ── Section 8: Compliance and Operations ─────────────────────────────────
    h1(doc, "8.  Compliance and Operations")

    h2(doc, "8.1  DNC Enforcement")
    para(doc,
        "Do-Not-Call compliance is enforced at two independent layers. This dual-layer design "
        "ensures that no DNC-registered number can be dialled regardless of application state "
        "or race conditions in the post-call pipeline."
    )
    h3(doc, "Layer 1 — Pre-Dial SQL Gate")
    para(doc,
        "The dialing worker fetches leads via a SQL query containing a NOT EXISTS subquery "
        "against dnc_registry. This gate operates at the database level and cannot be bypassed "
        "by application code:"
    )
    code(doc,
        "SELECT l.*\n"
        "FROM   agent_operations.leads l\n"
        "WHERE  l.status = 'pending'\n"
        "  AND  l.next_retry_at <= NOW()\n"
        "  AND  NOT EXISTS (\n"
        "         SELECT 1\n"
        "         FROM   agent_operations.dnc_registry d\n"
        "         WHERE  d.phone_number = l.phone_number\n"
        "       )\n"
        "ORDER BY l.created_at\n"
        "LIMIT 50;"
    )
    h3(doc, "Layer 2 — Post-Call Detection")
    para(doc,
        "The post-call analysis worker scans each transcript for DNC signals via two sub-layers: "
        "a Claude Sonnet context-aware classification (dnc_requested field in ExtractionResult) "
        "and a deterministic keyword scanner (dnc_keywords.scan()). Either layer triggering "
        "causes an immediate INSERT into dnc_registry (source='caller_request') and updates "
        "the lead status to 'failed_dnc'."
    )
    h3(doc, "Adding Numbers to DNC Manually")
    code(doc,
        "INSERT INTO agent_operations.dnc_registry (phone_number, source)\n"
        "VALUES ('+12125551234', 'manual');"
    )

    h2(doc, "8.2  TCPA Calling Hours")
    para(doc,
        "The dialing worker enforces calling hours of 08:00 to 21:00 in each lead's local "
        "timezone. The timezone value is an IANA timezone string (e.g. Asia/Kolkata, "
        "America/New_York) stored in leads.timezone at import time. The gate is evaluated "
        "per lead, enabling a single campaign to span multiple timezones correctly."
    )
    para(doc,
        "Leads outside calling hours are skipped in the current polling cycle and re-evaluated "
        "five seconds later. No lead status change is made for timezone-gated skips."
    )

    h2(doc, "8.3  KPI Targets and Alerting")
    tbl(doc,
        ["KPI", "Target", "Alert Threshold", "Data Source"],
        [
            ["End-to-end response latency", "< 500ms", "> 800ms sustained", "Retell AI analytics dashboard"],
            ["First-5-second AI detection rate", "< 20%", "> 35%", "ElevenLabs / Retell AI metrics"],
            ["Structured output completion rate", "> 95%", "< 85%", "call_transcripts.structured_data non-null count"],
            ["DNC scrub miss rate", "0%", "Any non-zero value", "kpi_check.py — live DB query"],
            ["Call abandon rate (< 10 seconds)", "< 15%", "> 25%", "call_logs.duration_sec histogram"],
            ["Average call duration", "> 90 seconds", "< 30 seconds sustained", "call_logs.duration_sec average"],
        ]
    )
    para(doc, "Evaluate current metrics against these thresholds:")
    code(doc, "python tests/evals/kpi_check.py")

    h2(doc, "8.4  Audit Trail")
    para(doc,
        "Every lead status change is timestamped via the leads.updated_at column, managed "
        "automatically by the SQLAlchemy TimestampMixin on every UPDATE operation. Call telemetry "
        "in call_logs is effectively immutable — the memories_role database role has no DELETE "
        "privilege on this table. Recording URLs stored in call_logs.recording_url should be "
        "reviewed and purged 90 days after the call unless the record has been flagged for "
        "retention."
    )

    h2(doc, "8.5  Known Gaps — Required Before Full Compliance")
    para(doc,
        "The following features are specified in the Product Requirements Document but have "
        "not yet been implemented. They are required for a fully compliant production deployment:"
    )
    tbl(doc,
        ["Gap", "Description", "Priority"],
        [
            ["GitHub Actions CI", "No .github/workflows/ directory exists. CI must run pytest and docker build on every pull request.", "High"],
            ["TCPA Litigator Scrub", "The PossibleNOW API (or equivalent) must be called at campaign activation to scrub leads against the national litigator registry.", "High"],
            ["Consent Timestamp Storage", "Lead import must record consent timestamp, disclosure text version, and source IP address alongside each lead record.", "High"],
            ["Recording Auto-Purge", "call_logs.recording_url entries must be purged 90 days post-call unless flagged. No automated purge job currently exists.", "Medium"],
            ["Latency Instrumentation", "End-to-end latency and first-5-second detection KPIs require additional columns in call_logs and Retell webhook metadata.", "Medium"],
        ]
    )

    pb(doc)

    # ── Appendix A ────────────────────────────────────────────────────────────
    h1(doc, "Appendix A — Environment Variables Reference")
    tbl(doc,
        ["Variable", "Required", "Default", "Service"],
        [
            ["POSTGRES_USER", "Yes", "postgres", "Database"],
            ["POSTGRES_PASSWORD", "Yes", "password", "Database"],
            ["POSTGRES_DB", "Yes", "voice_agent", "Database"],
            ["DATABASE_URL", "Yes", "—", "Database"],
            ["SQLALCHEMY_ECHO", "No", "false", "Database"],
            ["POOL_SIZE", "No", "10", "Database"],
            ["MAX_OVERFLOW", "No", "20", "Database"],
            ["REDIS_URL", "No", "redis://localhost:6379/0", "Queue"],
            ["RETELL_API_KEY", "Yes", "—", "Retell AI"],
            ["RETELL_WEBHOOK_SECRET", "Yes", "—", "Retell AI"],
            ["TELNYX_API_KEY", "Yes", "—", "Telephony"],
            ["TELNYX_FROM_NUMBER", "Yes", "+10000000000", "Telephony"],
            ["ELEVENLABS_API_KEY", "Yes", "—", "TTS"],
            ["ELEVENLABS_VOICE_ID", "Yes", "—", "TTS"],
            ["OPENAI_API_KEY", "Yes", "—", "LLM"],
            ["ANTHROPIC_API_KEY", "Yes", "—", "LLM"],
            ["WEBHOOK_PORT", "No", "8000", "API"],
            ["WEBHOOK_BASE_URL", "Yes", "http://localhost:8000", "API"],
            ["N8N_WEBHOOK_URL", "No", "http://localhost:5678/webhook/post-call", "Automation"],
            ["N8N_WEBHOOK_SECRET", "No", "changeme", "Automation"],
            ["N8N_BASIC_AUTH_USER", "No", "admin", "Automation"],
            ["N8N_BASIC_AUTH_PASSWORD", "No", "changeme", "Automation"],
            ["NEXT_PUBLIC_API_URL", "No", "http://localhost:8000", "Dashboard"],
            ["NEXT_PUBLIC_WS_URL", "No", "ws://localhost:8000", "Dashboard"],
        ]
    )

    pb(doc)

    # ── Appendix B ────────────────────────────────────────────────────────────
    h1(doc, "Appendix B — Glossary")
    terms = [
        ("Agent (Retell AI)", "A configured AI persona on the Retell AI platform, combining LLM model settings, ElevenLabs voice selection, and a system prompt template."),
        ("CNAM", "Caller Name — the display name shown to the called party. Controlled via the Telnyx phone number configuration."),
        ("CPS", "Calls Per Second — the outbound call initiation rate. This system enforces a hard limit of 1 CPS via asyncio.sleep(1.0) to comply with Telnyx carrier requirements."),
        ("DNC", "Do Not Call — a registry of phone numbers that must not be contacted for telemarketing or outbound sales purposes. Maintained in the dnc_registry table."),
        ("Dynamic Variables", "Per-call data injected into the Retell AI agent at call creation time (e.g. first_name, company, custom_vars). Enables personalised conversations at scale."),
        ("E.164", "International telephone number format: +[country code][subscriber number], e.g. +12125551234. All phone numbers in the system must use this format."),
        ("ExtractionResult", "The Pydantic model representing the structured output of the Claude Sonnet post-call analysis. Contains 9 typed fields: call_outcome, callback_time, objections_raised, next_action, summary, sentiment, sentiment_reason, lead_temperature, and dnc_requested."),
        ("HMAC-SHA256", "Hash-based Message Authentication Code using the SHA-256 digest algorithm. Used to verify that inbound webhook requests originate authentically from Retell AI."),
        ("IANA Timezone", "A timezone identifier from the IANA Time Zone Database (e.g. Asia/Kolkata, America/New_York, Europe/London). Used for per-lead calling hour enforcement."),
        ("JSONB", "PostgreSQL binary JSON storage type. Supports efficient indexing and querying of nested JSON structures. Used for prompt_template, llm_config, custom_vars, and structured_data columns."),
        ("RQ", "Redis Queue — a Python background job processing library built on Redis. Used by the post-call analysis worker to process call_analyzed events asynchronously."),
        ("SIP", "Session Initiation Protocol — the telephony signalling protocol used by Telnyx to establish, manage, and terminate voice calls."),
        ("STT", "Speech-to-Text — real-time transcription of spoken audio to text. Provided by Deepgram Nova-3 integrated into the Retell AI platform."),
        ("TCPA", "Telephone Consumer Protection Act — United States federal law governing telemarketing calls, requiring compliance with calling hours restrictions and Do Not Call registry obligations."),
        ("TTS", "Text-to-Speech — synthesis of natural-sounding spoken audio from text. Provided by ElevenLabs v3 with approximately 75ms first-audio latency."),
        ("VAD", "Voice Activity Detection — the process of detecting the onset and offset of human speech in an audio stream. Implemented using Silero VAD with a 200ms onset threshold and 800ms offset threshold."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{term}:  ")
        r1.font.name = FONT_BODY; r1.font.size = Pt(11); r1.bold = True
        r2 = p.add_run(definition)
        r2.font.name = FONT_BODY; r2.font.size = Pt(11)

    doc.save(output_path)
    print(f"Document written: {output_path}")


if __name__ == "__main__":
    build("voice-outbound-agent-documentation.docx")
